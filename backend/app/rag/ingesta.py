"""Utilidades compartidas de RAG: chunking, embeddings via Ollama, e indices Chroma.

Usado tanto por el RAG administrativo como por el RAG clinico (una coleccion
Chroma por paciente). La logica de construccion/consulta de indices vive aca
para no duplicarla entre ambos.
"""
from pathlib import Path
from typing import Any, Optional

import chromadb
from chromadb.api.types import Documents, EmbeddingFunction, Embeddings
from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from app.logging_config import get_logger
from app.ollama_client import embed

logger = get_logger("rag.ingesta")

# Chunking aproximado por caracteres (no hay tokenizer de Qwen a mano):
# ~4 caracteres por token, 400 tokens ~= 1600 caracteres, overlap de 50 tokens ~= 200 caracteres.
CHUNK_SIZE_CHARS = 1600
CHUNK_OVERLAP_CHARS = 200


class OllamaEmbeddingFunction(EmbeddingFunction):
    """Embedding function de Chroma que delega en el modelo de embeddings de Ollama."""

    def __init__(self) -> None:
        pass

    def __call__(self, input: Documents) -> Embeddings:
        return [embed(texto) for texto in input]

    def name(self) -> str:
        return "ollama_nomic_embed_text"

    def get_config(self) -> dict[str, Any]:
        return {}

    @staticmethod
    def build_from_config(config: dict[str, Any]) -> "OllamaEmbeddingFunction":
        return OllamaEmbeddingFunction()


def chunk_text(texto: str, chunk_size: int = CHUNK_SIZE_CHARS, overlap: int = CHUNK_OVERLAP_CHARS) -> list[str]:
    """Divide un texto en fragmentos con solapamiento, respetando saltos de linea
    cuando es posible para no cortar oraciones a la mitad.
    """
    texto = texto.strip()
    if len(texto) <= chunk_size:
        return [texto] if texto else []

    fragmentos = []
    inicio = 0
    while inicio < len(texto):
        fin = min(inicio + chunk_size, len(texto))
        if fin < len(texto):
            corte = texto.rfind("\n", inicio, fin)
            if corte == -1 or corte <= inicio:
                corte = texto.rfind(" ", inicio, fin)
            if corte > inicio:
                fin = corte
        fragmento = texto[inicio:fin].strip()
        if fragmento:
            fragmentos.append(fragmento)
        if fin >= len(texto):
            break
        inicio = max(fin - overlap, inicio + 1)

    return fragmentos


def _iterar_bloques_en_orden(documento: DocxDocument):
    """document.paragraphs y document.tables son listas separadas en python-docx
    y no preservan el orden real del documento (todos los parrafos primero, todas
    las tablas despues). Eso desconecta un encabezado de su tabla (ej. "Horario
    por especialidad (sede San Isidro)" terminaba lejos de la fila de Pediatria,
    y el LLM le atribuia la sede equivocada). Se recorre el XML del cuerpo en su
    orden real para que cada tabla quede junto al texto que la precede."""
    cuerpo = documento.element.body
    for hijo in cuerpo.iterchildren():
        if isinstance(hijo, CT_P):
            yield DocxParagraph(hijo, documento)
        elif isinstance(hijo, CT_Tbl):
            yield DocxTable(hijo, documento)


def _extraer_texto_docx(path: Path) -> str:
    """Extrae parrafos y tablas de un .docx como texto plano, en el orden real
    del documento. Las filas de tabla se serializan como 'celda | celda | ...'
    para conservar la relacion campo/valor."""
    documento = DocxDocument(str(path))
    partes = []
    for bloque in _iterar_bloques_en_orden(documento):
        if isinstance(bloque, DocxParagraph):
            if bloque.text.strip():
                partes.append(bloque.text.strip())
        else:  # DocxTable
            for fila in bloque.rows:
                celdas = [c.text.strip() for c in fila.cells]
                if any(celdas):
                    partes.append(" | ".join(celdas))
    return "\n".join(partes)


def _leer_documento(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return _extraer_texto_docx(path)
    return path.read_text(encoding="utf-8")


def _get_client(persist_dir: Path) -> chromadb.ClientAPI:
    persist_dir.mkdir(parents=True, exist_ok=True)
    return chromadb.PersistentClient(path=str(persist_dir))


def build_index(collection_name: str, persist_dir: Path, docs_dir: Path) -> int:
    """Reconstruye una coleccion desde cero a partir de los archivos .md/.txt/.docx
    de docs_dir.

    Retorna la cantidad de fragmentos indexados. Si docs_dir no tiene archivos,
    deja la coleccion vacia (las consultas responderan que no hay informacion).
    """
    client = _get_client(persist_dir)
    try:
        client.delete_collection(collection_name)
    except Exception:
        pass  # La coleccion no existia todavia, no hay nada que borrar.

    coleccion = client.create_collection(
        name=collection_name,
        embedding_function=OllamaEmbeddingFunction(),
        metadata={"hnsw:space": "cosine"},
    )

    # Se excluye README.md (guia de formato, no contenido real) y los archivos
    # temporales de Word (~$...) que quedan cuando un .docx esta abierto.
    archivos = (
        sorted(
            f
            for f in [*docs_dir.glob("*.md"), *docs_dir.glob("*.txt"), *docs_dir.glob("*.docx")]
            if f.name.lower() != "readme.md" and not f.name.startswith("~$")
        )
        if docs_dir.exists()
        else []
    )
    if not archivos:
        logger.warning("No hay documentos en %s, la coleccion '%s' queda vacia.", docs_dir, collection_name)
        return 0

    ids, documentos, metadatas = [], [], []
    for archivo in archivos:
        contenido = _leer_documento(archivo)
        for i, fragmento in enumerate(chunk_text(contenido)):
            ids.append(f"{archivo.stem}_{i}")
            documentos.append(fragmento)
            metadatas.append({"fuente": archivo.name})

    if documentos:
        coleccion.add(ids=ids, documents=documentos, metadatas=metadatas)

    logger.info("Indexados %d fragmentos de %d archivo(s) en '%s'.", len(documentos), len(archivos), collection_name)
    return len(documentos)


def query_index(
    collection_name: str,
    persist_dir: Path,
    query: str,
    top_k: int,
    min_similarity: float,
) -> list[dict]:
    """Recupera los top_k fragmentos mas similares a query, filtrando por similitud minima.

    Retorna lista vacia si la coleccion no existe, esta vacia, o ningun fragmento
    supera el umbral (el llamador debe responder que no tiene esa informacion).
    """
    client = _get_client(persist_dir)
    try:
        coleccion = client.get_collection(collection_name, embedding_function=OllamaEmbeddingFunction())
    except Exception:
        logger.warning("La coleccion '%s' no existe todavia en %s.", collection_name, persist_dir)
        return []

    if coleccion.count() == 0:
        return []

    resultado = coleccion.query(query_texts=[query], n_results=min(top_k, coleccion.count()))

    fragmentos = []
    documentos = resultado.get("documents", [[]])[0]
    metadatas = resultado.get("metadatas", [[]])[0]
    distancias = resultado.get("distances", [[]])[0]

    for doc, meta, distancia in zip(documentos, metadatas, distancias):
        similitud = 1 - distancia  # distancia coseno -> similitud coseno
        if similitud >= min_similarity:
            fragmentos.append({"texto": doc, "fuente": meta.get("fuente", "desconocido"), "similitud": similitud})

    return fragmentos
